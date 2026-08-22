from __future__ import annotations

import io

from werkzeug.datastructures import FileStorage

from phase1.parse import extract_upload, propose_from_text


def test_extract_docx_then_propose_django_not_k8s():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jordan Hale")
    doc.add_paragraph("jordan.hale@example.com")
    doc.add_paragraph("Built Django REST APIs on PostgreSQL and Docker at NimbusPay in 2024.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fs = FileStorage(stream=buf, filename="cv.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    text = extract_upload(fs)
    assert "Django" in text
    draft = propose_from_text(text)
    skills = {s.lower() for s in draft["skills"]}
    assert "django" in skills
    assert "kubernetes" not in skills


def test_markdown_to_pdf_is_a_pdf():
    from phase1.pack_files import markdown_to_pdf

    data = markdown_to_pdf("Ada Okafor\n\nPROFESSIONAL SUMMARY\nBackend engineer.\n")
    assert data[:4] == b"%PDF"
    assert b"Ada" in data or b"Okafor" in data or len(data) > 200


def test_extract_pdf_contains_python():
    raw = (
        b"%PDF-1.4\n1 0 obj<<>>endobj\n"
        b"stream\nBT (Python Django PostgreSQL) Tj ET\nendstream\n"
        b"%%EOF\n"
    )
    fs = FileStorage(stream=io.BytesIO(raw), filename="cv.pdf", content_type="application/pdf")
    text = extract_upload(fs)
    assert "Python" in text
    assert "Django" in text

from __future__ import annotations

import io
import re
from datetime import date

from phase0.truth import SKILL_CATALOG

_EMAIL = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
_YEAR = re.compile(r"\b(20\d{2})\b")
_PHONE = re.compile(r"(?:\+\d{1,3}[\s-]?)?(?:\(?\d{3,4}\)?[\s.-]?)\d{3}[\s.-]?\d{3,4}")
_MONTHS = r"jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?"
_DATE_SPAN = re.compile(
    rf"(?P<a>(?:(?:{_MONTHS})\.?\s+)?(?:19|20)\d{{2}})"
    rf"\s*[–\-—to]+\s*"
    rf"(?P<b>present|current|now|(?:(?:{_MONTHS})\.?\s+)?(?:19|20)\d{{2}})",
    re.I,
)
_DEGREE = re.compile(
    r"\b(?:b\.?\s*sc|b\.?\s*eng|b\.?\s*tech|b\.?\s*a\b|m\.?\s*sc|m\.?\s*eng|"
    r"mba|m\.?ba|hnd|ond|\bnd\b|ph\.?d|bachelor|master'?s|diploma|"
    r"ssce|waec|neco|\bnce\b|ll\.?b|ll\.?m)\b",
    re.I,
)
_HEAD_PATTERNS = (
    (
        "sum",
        r"professional\s+summary|career\s+summary|personal\s+profile|"
        r"career\s+objective|about\s+me|^summary$|^profile$|^objective$",
    ),
    (
        "skills",
        r"core\s+competenc|key\s+skills|professional\s+skills|technical\s+skills|"
        r"areas?\s+of\s+expertise|tech(?:nical)?\s+stack|^skills$|^technologies$|^tools$",
    ),
    (
        "edu",
        r"educational\s+qualif|educational\s+background|education\s+and\s+training|"
        r"education\s+and\s+cert|academic\s+qualif|^education$|^academic$|^qualifications$",
    ),
    ("cert", r"^certifications?$|^professional\s+certifications?$|^licen[cs]es?(?:\s+and\s+certifications?)?$"),
    (
        "exp",
        r"work\s+(?:history|experience)|professional\s+experience|employment\s+history|"
        r"^experience$|^employment$",
    ),
)
_TITLE_HINT = re.compile(
    r"\b(engineer|developer|analyst|manager|administrator|designer|scientist|"
    r"architect|specialist|consultant|intern|officer|lead|head of|director)\b",
    re.I,
)
_INLINE_SKILLS = re.compile(
    r"^(?:technical\s+)?(?:skills|technologies|tech stack|tools|competenc(?:e|ies|y)?)"
    r"\s*[:\-–]\s*(.+)$",
    re.I,
)
_SKILL_STOP = {
    "and",
    "or",
    "with",
    "including",
    "skills",
    "proficient",
    "expertise",
    "knowledge",
    "tools",
    "technologies",
    "stack",
    "others",
    "etc",
    "various",
}


def _skills_in(text: str) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()
    for skill in SKILL_CATALOG:
        if re.search(r"(?<![A-Za-z0-9])" + re.escape(skill) + r"(?![A-Za-z0-9])", text, re.I):
            key = skill.lower()
            if key in seen:
                continue
            seen.add(key)
            found.append(skill)
    return found


def _split_skill_tokens(blob: str) -> list[str]:
    blob = (blob or "").replace("•", ",").replace("|", ",").replace("·", ",").replace("/", ",")
    blob = re.sub(r"\s{2,}", ",", blob)
    blob = re.sub(r"[()]", " ", blob)
    parts = re.split(r"[,;]+", blob)
    out: list[str] = []
    seen: set[str] = set()
    for part in parts:
        token = re.sub(r"^[\-–—*]\s*", "", part).strip()
        token = re.sub(r"\s+", " ", token)
        if not token or token.lower() in _SKILL_STOP:
            continue
        if len(token) > 42:
            for hit in _skills_in(token):
                if hit.lower() not in seen:
                    seen.add(hit.lower())
                    out.append(hit)
            continue
        if len(token) < 2 and token.upper() not in {"C", "R", "GO"}:
            continue
        key = token.lower()
        if key not in seen:
            seen.add(key)
            out.append(token)
    if not out:
        out.extend(_skills_in(blob))
    return out


def _clean_profile_text(text: str) -> str:
    value = (text or "").strip()
    if not value:
        return ""
    value = _EMAIL.sub("", value)
    value = _PHONE.sub("", value)
    value = re.sub(r"\s+", " ", value).strip(" ,;|-")
    from phase0.geo import looks_like_place

    if looks_like_place(value) and len(value) < 50:
        return ""
    return value[:400]


def confirm_view(draft: dict) -> dict:
    """Values bound to Confirm — each box gets one kind of data."""
    from phase0.geo import auth_location_display, looks_like_place

    data = dict(draft or {})
    data["summary"] = _clean_profile_text(data.get("summary") or "")
    loc = auth_location_display(data)
    if loc and not looks_like_place(loc):
        loc = ", ".join(
            x for x in (data.get("work_authorization") or []) if x and str(x).upper() != "ANY"
        )
    data["auth_location"] = loc
    data["skills_text"] = ", ".join(data.get("skills") or [])
    data["education_text"] = "\n".join(data.get("education") or [])
    data["certifications_text"] = "\n".join(data.get("certifications") or [])
    return data


def merge_skills(*groups: list[str]) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for group in groups:
        for skill in group or []:
            key = " ".join(skill.lower().split())
            if not key or key in seen:
                continue
            seen.add(key)
            out.append(skill.strip())
    return out


def normalize_career_start(raw: str, fallback: str = "2023-01-01") -> str:
    value = (raw or "").strip()
    if re.fullmatch(r"\d{4}", value):
        value = f"{value}-01-01"
    elif re.fullmatch(r"\d{4}-\d{2}", value):
        value = f"{value}-01"
    try:
        parsed = date.fromisoformat(value[:10])
    except ValueError:
        return fallback
    this_year = date.today().year
    if parsed.year < 1990 or parsed.year > this_year:
        return fallback
    return parsed.isoformat()


def guess_work_authorization(text: str) -> list[str]:
    from phase0.geo import country_label, split_auth_location

    head = "\n".join((text or "").splitlines()[:12])
    labels, _loc = split_auth_location(head)
    blob = head.lower()
    digits = re.sub(r"\s+", "", head)
    if "+234" in digits or re.search(r"(?<!\d)234\d{7,}", digits):
        ng = country_label("NG")
        if ng not in labels:
            labels = [ng] + labels
    if labels:
        return labels
    if any(k in blob for k in ("nigeria", "lagos", "abuja", "port harcourt", "naija")):
        return [country_label("NG")]
    return []


def split_optional_lines(raw: str) -> list[str]:
    text = (raw or "").replace("\r\n", "\n").strip()
    if not text:
        return []
    parts = [ln.strip(" •-\t") for ln in text.split("\n")]
    return [p for p in parts if p]


def grounded_skills(listed: list[str], bullets: list[dict]) -> list[str]:
    from phase0.models import canon

    evidenced: set[str] = set()
    for b in bullets:
        for tag in b.get("tags") or []:
            evidenced.add(canon(tag))
        evidenced.update(canon(s) for s in _skills_in(b.get("text") or ""))
    keep = []
    seen = set()
    for s in listed:
        c = canon(s)
        if c in evidenced and c not in seen:
            seen.add(c)
            keep.append(s)
    return keep


def extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract_pdf(data: bytes) -> str:
    text = ""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        chunks: list[str] = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        text = "\n".join(chunks)
    except Exception:
        text = ""
    if len(text.strip()) < 40:
        strings = re.findall(rb"\(([^)]{3,})\)", data)
        extra = "\n".join(s.decode("latin-1", "replace") for s in strings)
        text = f"{text}\n{extra}".strip()
    return text


def _normalize_resume_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    text = _EMAIL.sub(lambda m: m.group(0) + "\n", text)
    text = _PHONE.sub(lambda m: "\n" + m.group(0) + "\n", text)
    text = re.sub(
        r"(?i)(?<=\S)\s+(professional\s+summary|personal\s+profile|career\s+objective|"
        r"core\s+competenc\w*|educational\s+qualif\w*|key\s+skills|work\s+experience|"
        r"education|experience|skills|certifications?)\b",
        r"\n\1",
        text,
    )
    return re.sub(r"\n{3,}", "\n\n", text)


def _kind_from_head(head: str) -> str | None:
    compact = re.sub(r"\s+", " ", (head or "").strip().lower())
    compact = compact.strip(":-– ")
    for kind, pattern in _HEAD_PATTERNS:
        if re.search(pattern, compact, re.I) or re.fullmatch(pattern, compact, re.I):
            return kind
    return None


def _split_section_line(ln: str) -> tuple[str, str] | None:
    raw = (ln or "").strip()
    if not raw:
        return None
    kind = _kind_from_head(raw)
    if kind:
        return kind, ""
    m = re.match(r"^(.{2,40}?)\s*[:\-–]\s*(.*)$", raw)
    if m:
        kind = _kind_from_head(m.group(1))
        if kind:
            return kind, m.group(2).strip()
    m = re.match(r"^(.{3,40}?)\s{2,}(.*)$", raw)
    if m:
        kind = _kind_from_head(m.group(1))
        if kind and kind != "exp":
            return kind, m.group(2).strip()
    parts = re.split(r"\s+", raw, maxsplit=3)
    for take in (3, 2, 1):
        if len(parts) <= take:
            continue
        head = " ".join(parts[:take])
        rest = " ".join(parts[take:]).strip()
        kind = _kind_from_head(head)
        if kind and kind != "exp" and rest and not _looks_like_role_header(rest):
            return kind, rest
    return None


def _looks_like_place_line(ln: str) -> bool:
    from phase0.geo import looks_like_place

    return looks_like_place(ln)


def _looks_like_role_header(ln: str) -> bool:
    if len(ln) > 140:
        return False
    if _DATE_SPAN.search(ln) and _TITLE_HINT.search(ln):
        return True
    if _DATE_SPAN.search(ln) and re.search(r"[,|@]| at ", ln):
        return True
    if _TITLE_HINT.search(ln) and re.search(r"\s(at|,|\||—|-)\s", ln) and len(ln) < 100:
        return True
    return False


def _split_role_header(ln: str) -> tuple[str, str, str, str]:
    start, end = "", "present"
    span = _DATE_SPAN.search(ln)
    rest = ln
    if span:
        start = re.sub(r"\s+", " ", span.group("a")).strip()
        end = re.sub(r"\s+", " ", span.group("b")).strip()
        rest = (ln[: span.start()] + " " + ln[span.end() :]).strip(" ,|–—-")
    rest = re.sub(r"[()]", " ", rest)
    rest = re.sub(r"\s+", " ", rest).strip(" ,|-")
    at = re.search(r"\bat\s+(.+)$", rest, re.I)
    if at:
        title = rest[: at.start()].strip(" ,|-") or "Role"
        employer = at.group(1).strip(" ,|-") or "Employer"
        return title[:80], employer[:80], start or "—", end
    parts = [p.strip() for p in re.split(r"\s*[|,]\s*", rest) if p.strip()]
    title = parts[0] if parts else "Role"
    employer = parts[1] if len(parts) > 1 else (parts[0] if parts else "Employer")
    if len(parts) == 1:
        employer = "Employer"
    return title[:80], employer[:80], start or "—", end


def _clean_bullet(ln: str) -> str:
    return re.sub(r"^[\-•–*—]\s*", "", ln).strip()


_SCHOOL = re.compile(
    r"university|polytechnic|college|secondary|high school|waec|neco|ssce|"
    r"school certificate|senior school|\bnce\b|\bhnd\b|\bond\b|bachelor|master|"
    r"doctorate|ph\.?d|b\.?\s*sc|m\.?\s*sc|ll\.?b",
    re.I,
)
_PRO_CERT = re.compile(
    r"\b(ccna|ccnp|ccie|pmp|itil|aws|azure|gcp|comptia|ceh|cissp|prince2|"
    r"scrum master|itil|microsoft certified|google certified|certified|"
    r"certification)\b",
    re.I,
)


def classify_credential(line: str) -> str | None:
    text = (line or "").strip()
    if not text:
        return None
    schoolish = bool(_DEGREE.search(text) or _SCHOOL.search(text))
    certish = bool(_PRO_CERT.search(text))
    if schoolish and re.search(r"school certificate|senior school|waec|neco|ssce", text, re.I):
        return "edu"
    if schoolish and not certish:
        return "edu"
    if certish and not schoolish:
        return "cert"
    if schoolish:
        return "edu"
    return None


def partition_credentials(education: list[str], certifications: list[str]) -> tuple[list[str], list[str]]:
    edu: list[str] = []
    cert: list[str] = []
    seen: set[str] = set()
    for origin, bucket in ((education, "edu"), (certifications, "cert")):
        for line in origin or []:
            key = re.sub(r"\s+", " ", line).strip().lower()
            if not key or key in seen:
                continue
            seen.add(key)
            kind = classify_credential(line) or bucket
            if kind == "edu":
                edu.append(line)
            else:
                cert.append(line)
    return edu, cert


def extract_upload(file_storage) -> str:
    name = (getattr(file_storage, "filename", None) or "").lower()
    data = file_storage.read()
    if hasattr(file_storage, "seek"):
        try:
            file_storage.seek(0)
        except Exception:
            pass
    if name.endswith(".docx"):
        return extract_docx(data)
    if name.endswith(".pdf"):
        return extract_pdf(data)
    if name.endswith(".txt") or name.endswith(".md") or name.endswith(".rtf"):
        return data.decode("utf-8", "replace")
    raise ValueError("Upload a PDF or DOCX, or paste the résumé text.")


def propose_from_text(text: str) -> dict:
    text = _normalize_resume_text(text)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    email_m = _EMAIL.search(text)
    email = email_m.group(0) if email_m else ""
    name = lines[0] if lines and "@" not in lines[0] else "Candidate"
    location = ""
    for ln in lines[:12]:
        if ln == name or "@" in ln:
            continue
        if _PHONE.search(ln) and len(ln) < 28:
            continue
        if _split_section_line(ln):
            break
        if _looks_like_place_line(ln) and not _looks_like_role_header(ln):
            location = ln
            break
    phone = ""
    ph = _PHONE.search(text)
    if ph:
        phone = ph.group(0).strip()
    this_year = date.today().year
    years = [int(y) for y in _YEAR.findall(text) if 1995 <= int(y) <= this_year]
    career_start = normalize_career_start(f"{min(years)}-01-01" if years else "2023-01-01")  # replaced after roles
    catalog_skills = _skills_in(text)
    section_skills: list[str] = []

    section = ""
    roles: list[dict] = []
    education: list[str] = []
    certifications: list[str] = []
    summary_bits: list[str] = []
    loose: list[dict] = []
    current: dict | None = None

    def flush() -> None:
        nonlocal current
        if current is None:
            return
        if current["bullets"] or current["title"] not in {"Role"}:
            roles.append(current)
        current = None

    for ln in lines:
        if ln in {name, email, phone} or (location and ln == location):
            continue
        inline_skills = _INLINE_SKILLS.match(ln)
        if inline_skills:
            section_skills.extend(_split_skill_tokens(inline_skills.group(1)))
            continue
        split = _split_section_line(ln)
        if split:
            key, rest = split
            flush()
            section = key
            if rest:
                if section == "edu":
                    education.append(rest)
                elif section == "cert":
                    certifications.append(_clean_bullet(rest))
                elif section == "sum":
                    summary_bits.append(_clean_bullet(rest))
                elif section == "skills":
                    section_skills.extend(_split_skill_tokens(rest))
            continue
        if section == "edu":
            token = _clean_bullet(ln)
            kind = classify_credential(token)
            if kind == "cert" and token:
                certifications.append(token)
            elif token and (kind == "edu" or len(token) > 8):
                education.append(token)
            continue
        if section == "cert":
            token = _clean_bullet(ln)
            kind = classify_credential(token)
            if kind == "edu" and token:
                education.append(token)
            elif len(token) >= 2:
                certifications.append(token)
            continue
        if section == "sum":
            if len(ln) > 20:
                summary_bits.append(_clean_bullet(ln))
            continue
        if section == "skills":
            section_skills.extend(_split_skill_tokens(ln))
            continue
        if _looks_like_role_header(ln):
            flush()
            title, employer, start, end = _split_role_header(ln)
            current = {
                "id": f"r.{len(roles)}",
                "title": title,
                "employer": employer,
                "start": start,
                "end": end,
                "bullets": [],
            }
            section = "exp"
            continue
        cleaned = _clean_bullet(ln)
        if re.match(r"(?i)^(skills|technologies|tech stack|tools|competenc)\b", cleaned):
            extra = _INLINE_SKILLS.match(cleaned)
            if extra:
                section_skills.extend(_split_skill_tokens(extra.group(1)))
            continue
        tags = _skills_in(cleaned)
        is_bullet = bool(re.match(r"^[\-•–*—]\s+", ln)) or len(cleaned) >= 28
        if current is not None and is_bullet:
            current["bullets"].append(
                {
                    "id": f"r.{len(roles)}.{len(current['bullets'])}",
                    "text": cleaned,
                    "tags": tags,
                }
            )
            continue
        if is_bullet and (tags or len(cleaned) >= 40):
            loose.append({"id": f"u.{len(loose)}", "text": cleaned, "tags": tags})

    flush()
    if not roles and loose:
        roles.append(
            {
                "id": "from-resume",
                "employer": "From résumé",
                "title": "Roles as stated",
                "start": career_start[:7],
                "end": "present",
                "bullets": loose,
            }
        )
    bullets = [b for role in roles for b in role["bullets"]]
    employers = []
    for role in roles:
        emp = role.get("employer") or ""
        if emp and emp not in employers and emp.lower() not in {"employer", "from résumé", "from resume"}:
            employers.append(emp)
    if not summary_bits:
        for ln in lines[1:14]:
            if ln in {name, email, phone, location}:
                continue
            if _EMAIL.search(ln) or (_PHONE.search(ln) and len(ln) < 28):
                continue
            if _looks_like_place_line(ln) or _looks_like_role_header(ln):
                continue
            if _split_section_line(ln):
                break
            if len(ln) >= 40:
                summary_bits.append(ln)
            if sum(len(x) for x in summary_bits) > 240:
                break
    if not education:
        for ln in lines:
            if _DEGREE.search(ln) and 10 < len(ln) < 160 and not _looks_like_role_header(ln):
                if ln not in education:
                    education.append(ln)
    summary = " ".join(summary_bits).strip()
    if not summary:
        for b in bullets[:2]:
            summary_bits.append(b["text"].rstrip("."))
        if summary_bits:
            summary = ". ".join(summary_bits) + "."
    role_years: list[int] = []
    for role in roles:
        role_years.extend(int(y) for y in _YEAR.findall(role.get("start") or "") if 1995 <= int(y) <= this_year)
    if role_years:
        career_start = normalize_career_start(f"{min(role_years)}-01-01")
    elif years:
        edu_years = {int(y) for row in education for y in _YEAR.findall(row)}
        work_years = [y for y in years if y not in edu_years] or years
        career_start = normalize_career_start(f"{min(work_years)}-01-01")
    auth = guess_work_authorization(text)
    if not location or location.lower() == "not specified":
        location = auth[0] if auth else ""
    summary = _clean_profile_text(summary)
    education, certifications = partition_credentials(education, certifications)
    return {
        "name": name[:80],
        "email": email,
        "location": location,
        "phone": phone,
        "remote_ok": True,
        "work_authorization": auth,
        "career_start": career_start,
        "skills": merge_skills(section_skills, catalog_skills, [t for b in bullets for t in (b.get("tags") or [])]),
        "employers": employers,
        "summary": summary[:400],
        "education": education[:8],
        "certifications": certifications[:12],
        "experience": roles,
    }
